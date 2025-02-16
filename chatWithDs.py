import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import asyncio
from dotenv import load_dotenv
from docx import Document
from PyPDF2 import PdfReader
from openai import OpenAI
import logging
import ctypes
from ctypes import sizeof, windll, byref, c_int
import threading
from concurrent.futures import ThreadPoolExecutor
from threading import Semaphore

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class ConversationManager:
    def __init__(self):
        self.history = []
        self.max_length = 15  # 保持最近15组对话

    def add_message(self, role, content):
        self.history.append({"role": role, "content": content})
        if len(self.history) > self.max_length * 2:
            self.history = self.history[-self.max_length * 2:]

class ChatbotGUI:
    def __init__(self):
        """Initialize the chatbot GUI"""
        self.root = tk.Tk()
        self.root.title("Chat with DeepSeek")
        
        # 获取屏幕宽度
        screen_width = self.root.winfo_screenwidth()
        
        # 设置窗口大小和位置
        window_height = int(screen_width * 0.4)  # 降低高度比例
        self.root.geometry(f"{screen_width}x{window_height}+0+0")
        
        # 设置颜色主题
        self.colors = {
            'bg': '#f0f0f0',  # 背景色
            'chat_bg': '#ffffff',  # 聊天区域背景色
            'input_bg': '#ffffff',  # 输入区域背景色
            'button_bg': '#FFB6C1',  # 按钮背景色（粉色）
            'button_fg': '#0000FF',  # 按钮文字颜色（蓝色）
            'input_fg': '#0000FF',  # 输入文字颜色（蓝色）
        }
        
        # 设置窗口背景色
        self.root.configure(bg=self.colors['bg'])
        
        # 设置默认字体
        self.root.option_add('*Font', ('Microsoft YaHei', 12))
        
        # 设置当前模式
        self.current_mode = tk.StringVar(value="Chat")
        
        # 初始化对话管理器
        self.conversation_manager = ConversationManager()
        
        # API Key
        self.api_key = tk.StringVar()  # 使用StringVar来存储API密钥
        
        self.check_and_create_env()
        self.prompt_for_api_key()
        
        # Initialize OpenAI client
        self.client = OpenAI(
            base_url="https://api.deepseek.com/v1",
            api_key=self.api_key.get()
        )
        
        self.api_semaphore = Semaphore(3)  # 限制并发请求数
        
        self.setup_ui()
        
        # Initialize thread pool
        self.thread_pool = ThreadPoolExecutor(max_workers=5)
        
    def setup_ui(self):
        # 定义字体
        default_font = ('Microsoft YaHei', 12)
        button_font = ('Microsoft YaHei', 16)  # 按钮使用大号字体
        text_font = ('Microsoft YaHei', 12)
        
        # Chat display area
        self.chat_frame = ttk.Frame(self.root, style='Chat.TFrame')
        self.chat_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=(5, 0))
        
        # 创建自定义样式
        style = ttk.Style()
        style.configure('Chat.TFrame', background=self.colors['bg'])
        style.configure('Custom.TButton',
                       background=self.colors['button_bg'],
                       foreground=self.colors['button_fg'],
                       padding=5,
                       font=button_font)  # 使用大号字体
        
        # 配置 Radiobutton 样式
        style.configure('Custom.TRadiobutton',
                       font=button_font,  # 使用大号字体
                       background=self.colors['bg'])
        
        self.chat_display = tk.Text(
            self.chat_frame,
            wrap=tk.WORD,
            state=tk.DISABLED,
            height=15,
            font=text_font,
            bg=self.colors['chat_bg'],
            fg=self.colors['input_fg'],
            padx=10,
            pady=10,
            relief=tk.SOLID,
            borderwidth=1
        )
        self.chat_display.pack(fill=tk.BOTH, expand=True)
        
        # Scrollbar
        scrollbar = ttk.Scrollbar(self.chat_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.chat_display.config(yscrollcommand=scrollbar.set)
        scrollbar.config(command=self.chat_display.yview)

        # 创建一个框架来容纳loading和按钮
        self.control_frame = ttk.Frame(self.root, style='Chat.TFrame')
        self.control_frame.pack(fill=tk.X, padx=10, pady=(0, 5))
        
        # Loading indicator with larger red font
        self.loading_label = ttk.Label(
            self.control_frame,
            text="加载中，请耐心等待loading...",
            font=('Microsoft YaHei', 14, 'bold'),
            foreground='#ff0000'
        )
        self.loading_label.pack(side=tk.LEFT, padx=(50, 0))  # 添加左侧padding
        self.loading_label.pack_forget()  # 初始时隐藏
        
        # 创建模式选择框架
        self.mode_frame = ttk.Frame(self.control_frame, style='Chat.TFrame')
        self.mode_frame.pack(side=tk.LEFT, padx=(10, 0))
        
        # 创建模式选择单选按钮
        self.chat_mode = ttk.Radiobutton(
            self.mode_frame,
            text="聊天模式",
            variable=self.current_mode,
            value="Chat",
            style='Custom.TRadiobutton'  # 使用自定义样式
        )
        self.chat_mode.pack(side=tk.LEFT, padx=(0, 5))
        
        self.reasoner_mode = ttk.Radiobutton(
            self.mode_frame,
            text="推理模式",
            variable=self.current_mode,
            value="Reasoner",
            style='Custom.TRadiobutton'  # 使用自定义样式
        )
        self.reasoner_mode.pack(side=tk.LEFT)
        
        # 创建一个框架来容纳按钮，使用place而不是pack
        self.button_container = ttk.Frame(self.control_frame, style='Chat.TFrame')
        self.button_container.pack(side=tk.LEFT, padx=(0, 5))
        
        # File upload button
        self.upload_btn = ttk.Button(
            self.button_container,
            text="上传文件",
            style='Custom.TButton',
            command=self.upload_file
        )
        self.upload_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        # Download button
        self.download_btn = ttk.Button(
            self.button_container,
            text="下载聊天记录",
            style='Custom.TButton',
            command=self.download_chat
        )
        self.download_btn.pack(side=tk.LEFT)
        
        # Input area
        self.input_frame = ttk.Frame(self.root, style='Chat.TFrame')
        self.input_frame.pack(fill=tk.X, padx=10, pady=5)
        
        self.input_box = tk.Text(
            self.input_frame,
            height=3,
            wrap=tk.WORD,
            font=text_font,
            bg=self.colors['input_bg'],
            fg=self.colors['input_fg'],
            relief=tk.SOLID,
            borderwidth=1,
            padx=10,
            pady=5,
            insertwidth=2,
            insertbackground=self.colors['input_fg']
        )
        self.input_box.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # 设置输入框焦点
        self.input_box.focus_set()
        
        # 绑定回车键事件
        self.input_box.bind("<Return>", self.handle_return)
        self.last_return_time = 0  # 用于追踪上一次回车的时间
        
        # 创建一个框架来容纳Send和Clear按钮
        self.button_frame = ttk.Frame(self.input_frame, style='Chat.TFrame')
        self.button_frame.pack(side=tk.RIGHT, padx=5)
        
        # Send button
        self.send_btn = ttk.Button(
            self.button_frame,
            text="发送",
            style='Custom.TButton',
            command=self.send_message_event
        )
        self.send_btn.pack(side=tk.LEFT, padx=2)
        
        # Clear button
        self.clear_btn = ttk.Button(
            self.button_frame,
            text="清除",
            style='Custom.TButton',
            command=self.clear_input
        )
        self.clear_btn.pack(side=tk.LEFT, padx=2)
        
        # Bind Ctrl+Return to send
        self.input_box.bind("<Control-Return>", self.send_message_event)
        
        # 添加底部空白区域
        self.bottom_frame = ttk.Frame(self.root, style='Chat.TFrame')
        self.bottom_frame.pack(fill=tk.X, pady=(0, 30))  # 增加底部边距
        
    def clear_input(self):
        """Clear the input box"""
        self.input_box.delete('1.0', tk.END)
        self.input_box.focus_set()  # 清除后重新获得焦点
        
    def read_docx(self, file_path):
        """Read content from Word document"""
        try:
            doc = Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            logging.error(f"Error reading docx file: {e}")
            messagebox.showerror("Error", f"Could not read docx file: {e}")
            return None
        
    def read_pdf(self, file_path):
        """Read content from PDF document"""
        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                return text
        except Exception as e:
            logging.error(f"Error reading pdf file: {e}")
            messagebox.showerror("Error", f"Could not read pdf file: {e}")
            return None
            
    def read_txt(self, file_path):
        """Read content from text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logging.error(f"Error reading txt file: {e}")
            messagebox.showerror("Error", f"Could not read txt file: {e}")
            return None
            
    ALLOWED_EXTENSIONS = {'.docx', '.pdf', '.txt'}
    
    def upload_file(self):
        """Handle file upload"""
        file_types = [
            ('Word Documents', '*.docx'),
            ('PDF Files', '*.pdf'),
            ('Text Files', '*.txt')
        ]
        
        file_path = filedialog.askopenfilename(
            title="Select File",
            filetypes=file_types
        )
        
        if not file_path:
            return
        
        if os.path.splitext(file_path)[1].lower() not in self.ALLOWED_EXTENSIONS:
            messagebox.showerror("Error", "Unsupported file type")
            return
        
        # Read content based on file type
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.docx':
            self.thread_pool.submit(self.process_file, file_path, ext)
        elif ext == '.pdf':
            self.thread_pool.submit(self.process_file, file_path, ext)
        elif ext == '.txt':
            self.thread_pool.submit(self.process_file, file_path, ext)
        else:
            messagebox.showerror("Error", "Unsupported file type")
            return
            
    def process_file(self, file_path, ext):
        """Process file in a separate thread"""
        if ext == '.docx':
            content = self.read_docx(file_path)
        elif ext == '.pdf':
            content = self.read_pdf(file_path)
        elif ext == '.txt':
            content = self.read_txt(file_path)
        else:
            return
        
        if content:
            # Insert content into input box
            self.root.after(0, self.insert_content, content)
            
            # Automatically send message
            self.root.after(0, self.send_message_event)
            
    def insert_content(self, content):
        """Insert content into input box"""
        self.input_box.delete('1.0', tk.END)
        self.input_box.insert('1.0', content)
        
    async def send_message(self, message):
        with self.api_semaphore:
            try:
                # 根据当前模式选择不同的模型和系统消息
                if self.current_mode.get() == "Chat":
                    model = "deepseek-chat"
                    system_message = ""
                else:  # Reasoner mode
                    model = "deepseek-reasoner"
                    system_message = ""
                
                # 构建完整的消息历史
                messages = self.conversation_manager.history.copy()
                if not any(msg["role"] == "system" for msg in messages):
                    messages.insert(0, {"role": "system", "content": system_message})
                messages.extend([
                    msg for msg in self.conversation_manager.history 
                    if msg["role"] in ("user", "assistant")
                ])
                messages.append({"role": "user", "content": message})
                
                response = await asyncio.to_thread(
                    self.client.chat.completions.create,
                    model=model,
                    messages=messages
                )
                
                logging.info(f"API Response: {response}")
                
                reply = response.choices[0].message.content
                
                # 更新对话历史
                self.conversation_manager.add_message("user", message)
                self.conversation_manager.add_message("assistant", reply)
                
                return reply
                
            except Exception as e:
                logging.error(f"API call failed: {e}")
                messagebox.showerror("Error", f"API call failed: {str(e)}")
                return None

    def display_message(self, message, is_user=True):
        """Display a message in the chat display"""
        prefix = "You: " if is_user else "Assistant: "
        tag_name = 'user_msg' if is_user else 'ai_msg'
        
        self.chat_display.config(state=tk.NORMAL)
        if self.chat_display.get('1.0', tk.END).strip():
            self.chat_display.insert(tk.END, '\n\n')
        
        # 创建带颜色的标签
        self.chat_display.tag_config('user_msg', foreground=self.colors['button_fg'])
        self.chat_display.tag_config('ai_msg', foreground=self.colors['button_fg'])
        
        # 插入带颜色的消息
        self.chat_display.insert(tk.END, prefix, tag_name)
        self.chat_display.insert(tk.END, message)
        
        self.chat_display.config(state=tk.DISABLED)
        self.chat_display.see(tk.END)

    def stream_display(self, message):
        """Display a message in the chat display"""
        prefix = "Assistant: "
        tag_name = 'ai_msg'
        
        self.chat_display.config(state=tk.NORMAL)
        if self.chat_display.get('1.0', tk.END).strip():
            self.chat_display.insert(tk.END, '\n')
        
        # 创建带颜色的标签
        self.chat_display.tag_config('ai_msg', foreground=self.colors['button_fg'])
        
        # 插入带颜色的消息
        self.chat_display.insert(tk.END, prefix, tag_name)
        self.chat_display.insert(tk.END, message)
        
        self.chat_display.config(state=tk.DISABLED)
        self.chat_display.see(tk.END)

    def send_message_event(self, event=None):
        """Event handler for send message"""
        message = self.input_box.get('1.0', tk.END).strip()
        if not message:
            return

        # Display user message
        self.display_message(message, is_user=True)
        
        # Clear input box
        self.input_box.delete('1.0', tk.END)
        self.input_box.focus_set()  # 发送后重新获得焦点
        
        # Show loading indicator
        self.loading_label.pack()
        self.root.update()

        # Create a new thread to handle the async operation
        def async_handler():
            async def async_operation():
                try:
                    # Get AI response
                    ai_response = await self.send_message(message)
                    if ai_response:
                        # Schedule the UI update in the main thread
                        self.root.after(0, lambda: self.display_message(ai_response, is_user=False))
                finally:
                    # Schedule hiding the loading indicator in the main thread
                    self.root.after(0, lambda: self.loading_label.pack_forget())
                    self.root.after(0, self.root.update)

            # Create and run the event loop
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                loop.run_until_complete(async_operation())
            finally:
                loop.close()

        # Start the async handler in a separate thread
        threading.Thread(target=async_handler, daemon=True).start()

    def handle_return(self, event):
        """智能处理回车键事件"""
        current_time = event.time
        
        # 获取当前光标位置
        cursor_pos = self.input_box.index(tk.INSERT)
        line_start = self.input_box.index(f"{cursor_pos} linestart")
        line_end = self.input_box.index(f"{cursor_pos} lineend")
        current_line = self.input_box.get(line_start, line_end)
        
        # 如果是在空行上按回车
        if not current_line.strip():
            # 检查是否是双击回车（两次回车间隔小于300毫秒）
            if current_time - self.last_return_time < 300:  
                # 删除多余的空行
                self.input_box.delete(line_start, line_end + "+1c")
                # 发送消息
                self.send_message_event()
                return "break"  # 阻止默认的回车行为
        
        # 更新上一次回车时间
        self.last_return_time = current_time
        
        # 允许正常的回车换行
        return None

    def download_chat(self):
        """将聊天内容下载为Word文档"""
        try:
            # 获取保存文件的路径
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Document", "*.docx")],
                title="Save Chat As"
            )
            
            if not file_path:  # 如果用户取消了保存对话框
                return
                
            # 创建新的Word文档
            doc = Document()
            
            # 设置默认字体为雅黑
            style = doc.styles['Normal']
            style.font.name = 'Microsoft YaHei'
            
            # 添加标题
            heading = doc.add_heading('Chat History', 0)
            heading.style.font.name = 'Microsoft YaHei'  # 标题也使用雅黑字体
            
            # 获取聊天内容
            chat_content = self.chat_display.get('1.0', tk.END).strip()
            
            # 如果聊天内容为空
            if not chat_content:
                messagebox.showinfo("Info", "No chat content to download.")
                return
                
            # 按行分割聊天内容
            lines = chat_content.split('\n')
            current_paragraph = []
            
            # 处理每一行
            for line in lines:
                if line.strip():  # 如果不是空行
                    if line.startswith(("You:", "Assistant:")):  # 新消息开始
                        # 如果有累积的段落内容，先添加它
                        if current_paragraph:
                            p = doc.add_paragraph()
                            p.add_run(''.join(current_paragraph)).font.name = 'Microsoft YaHei'
                            current_paragraph = []
                        # 添加新的发言者标记
                        p = doc.add_paragraph()
                        run = p.add_run(line)
                        run.bold = True
                        run.font.name = 'Microsoft YaHei'
                    else:
                        current_paragraph.append(line + '\n')
            
            # 添加最后一个段落
            if current_paragraph:
                p = doc.add_paragraph()
                p.add_run(''.join(current_paragraph)).font.name = 'Microsoft YaHei'
            
            # 保存文档
            doc.save(file_path)
            messagebox.showinfo("Success", "Chat history has been saved successfully!")
            
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred while saving: {str(e)}")
            
    def check_and_create_env(self):
        if not os.path.exists('.env'):
            with open('.env', 'w') as f:
                f.write('DEEPSEEK_API_KEY=\n')  # 创建空的API密钥

    def prompt_for_api_key(self):
        self.check_and_create_env()
        api_key_window = tk.Toplevel(self.root)
        api_key_window.title("Enter API Key")
        api_key_label = tk.Label(api_key_window, text="Enter your API Key:")
        api_key_label.pack()
        api_key_entry = tk.Entry(api_key_window, textvariable=self.api_key)
        api_key_entry.pack()
        def save_api_key():
            api_key = self.api_key.get()  # 获取输入的API密钥
            if api_key:
                with open('.env', 'a') as f:
                    f.write(f'DEEPSEEK_API_KEY={api_key}\n')
                messagebox.showinfo('Success', 'API Key saved successfully!')
                api_key_window.destroy()
            else:
                messagebox.showerror('Error', 'Please enter a valid API Key.')
        save_button = tk.Button(api_key_window, text="Save", command=save_api_key)
        save_button.pack()

    def run(self):
        try:
            self.root.mainloop()
        except Exception as e:
            logging.error(f'An error occurred: {e}')
            messagebox.showerror('Error', 'An unexpected error occurred. Please check the logs for details.')
            self.root.destroy()

def get_api_key():
    with open('config.txt', 'r') as file:
        key_prompt = file.readline().strip()
        api_key = input(key_prompt)  # 提示用户输入 API Key
    return api_key

if __name__ == "__main__":
    try:
        app = ChatbotGUI()
        app.run()
    except Exception as e:
        logging.error(f'An error occurred: {e}')
        messagebox.showerror('Error', 'An unexpected error occurred. Please check the logs for details.')